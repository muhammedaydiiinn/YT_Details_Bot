import requests
import re
from docx import Document


def get_video_id(url):
    # YouTube URL'sinden video ID'sini çıkar
    pattern = r"(?<=v=)[a-zA-Z0-9_-]+(?=&|\?|$)"
    match = re.search(pattern, url)
    if match:
        return match.group()
    else:
        return None


def get_video_details(video_url, api_key):
    video_id = get_video_id(video_url)
    if video_id:
        url = f"https://www.googleapis.com/youtube/v3/videos?part=snippet&id={video_id}&key={api_key}"
        response = requests.get(url)
        data = response.json()

        if 'items' in data and len(data['items']) > 0:
            snippet = data['items'][0]['snippet']
            title = snippet['title']
            description = snippet['description']

            # Ayrıntıları döndür
            return title, description
        else:
            return "Video not found or details could not be retrieved.", ""
    else:
        return "Invalid YouTube URL.", ""


def save_video_details_to_docx(input_file, output_file):
    # Yeni bir Word belgesi oluştur
    doc = Document()

    with open(input_file, "r") as file:
        lines = file.readlines()
        api_key = lines[0].strip()

        # Video URL'yi çıkarmak için diğer satırları döngüye al
        for line in lines[1:]:
            video_url = line.strip()
            title, description = get_video_details(video_url, api_key)

            # Başlık ekle
            doc.add_heading(title, level=1)

            # Açıklama ekle
            doc.add_paragraph(description)

            # Boş satır ekle
            doc.add_paragraph()

    # Dosyayı kaydet
    doc.save(output_file)


# Kullanım örneği:
input_file = "info.txt"
output_file = "video_details.docx"
save_video_details_to_docx(input_file, output_file)
